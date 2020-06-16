from gtts import gTTS
import os
from kivy.uix.label import Label
from kivy.uix.popup import Popup
from kivy.uix.floatlayout import FloatLayout
from kivy.properties import ObjectProperty, ListProperty, StringProperty
from kivy.factory import Factory
from kivy.app import App
from kivy.config import Config
from tika import parser
from docx import Document
Config.set('graphics', 'resizable', False)


class LoadDialog(FloatLayout):
    load = ObjectProperty(None)
    cancel = ObjectProperty(None)
    defaultPath = os.path.join(
        os.environ['SYSTEMDRIVE'], '/Users', os.getlogin())


class LangDialog(FloatLayout):
    save = ObjectProperty(None)
    cancel = ObjectProperty(None)

    f = ''
    line = ''
    listValues = ListProperty()
    initialValue = StringProperty()
    listFiles = StringProperty()

    def __init__(self, **kwargs):
        super(LangDialog, self).__init__(**kwargs)
        self.listValues = ['af: Afrikaans', 'ar: Arabic', 'bn: Bengali', 'bs: Bosnian', 'ca: Catalan', 'cs: Czech', 'cy: Welsh', 'da: Danish', 'de: German', 'el: Greek', 'en-au: English (Australia)', 'en-ca: English (Canada)', 'en-gb: English (UK)', 'en-gh: English (Ghana)', 'en-ie: English (Ireland)', 'en-in: English (India)', 'en-ng: English (Nigeria)', 'en-nz: English (New Zealand)', 'en-ph: English (Philippines)', 'en-tz: English (Tanzania)', 'en-uk: English (UK)', 'en-us: English (US)', 'en-za: English (South Africa)', 'en: English', 'eo: Esperanto', 'es-es: Spanish (Spain)', 'es-us: Spanish (United States)', 'es: Spanish', 'et: Estonian', 'fi: Finnish', 'fr-ca: French (Canada)', 'fr-fr: French (France)', 'fr: French', 'gu: Gujarati', 'hi: Hindi', 'hr: Croatian', 'hu: Hungarian', 'hy: Armenian', 'id: Indonesian', 'is: Icelandic', 'it: Italian', 'ja: Japanese', 'jw: Javanese', 'km: Khmer', 'kn: Kannada', 'ko: Korean', 'la: Latin', 'lv: Latvian', 'mk: Macedonian', 'ml: Malayalam', 'mr: Marathi', 'my: Myanmar (Burmese)', 'ne: Nepali', 'nl: Dutch', 'no: Norwegian', 'pl: Polish', 'pt-br: Portuguese (Brazil)', 'pt-pt: Portuguese (Portugal)', 'pt: Portuguese', 'ro: Romanian', 'ru: Russian', 'si: Sinhala',
                           'sk: Slovak', 'sq: Albanian', 'sr: Serbian', 'su: Sundanese', 'sv: Swedish', 'sw: Swahili', 'ta: Tamil', 'te: Telugu', 'th: Thai', 'tl: Filipino', 'tr: Turkish', 'uk: Ukrainian', 'ur: Urdu', 'vi: Vietnamese', 'zh-cn: Chinese (Mandarin/China)', 'zh-tw: Chinese (Mandarin/Taiwan)']
        self.listFiles = '\n'.join(Root.loadedList)
        self.initialValue = "Clique aqui para escolher a língua"


class Root(FloatLayout):
    App.title = 'Conversor de Texto para MP3'
    loadfile = ObjectProperty(None)
    savefile = ObjectProperty(None)

    choosenLang = ''
    mimeTypes = ['.pdf', '.txt', '.docx']
    fullpath = []
    loadedList = []
    finishedList = []
    failedList = []

    def dismiss_popup(self):
        self._popup.dismiss()

    def show_load(self):
        content = LoadDialog(load=self.load, cancel=self.dismiss_popup)
        self._popup = Popup(title="Selecionar arquivo para conversão", content=content,
                            size_hint=(0.9, 0.9))
        self._popup.open()

    def load(self, path, filename):
        if(len(filename) != 0):
            mime = filename[0].split('.')
            mime = '.' + mime[len(mime) - 1]
            if (mime in Root.mimeTypes):
                separator = '\n'
                name = filename[0].split('\\')
                length = len(name) - 1
                Root.fullpath.append(filename[0])
                Root.loadedList.append(name[length])
                self.label1_wid.text = separator.join(Root.loadedList)
                self.dismiss_popup()

    def show_lang(self):
        if (len(Root.loadedList) == 0):
            self.label1_wid.text = 'Não há arquivos na fila'
        else:
            content = LangDialog(save=self.save, cancel=self.dismiss_popup)
            self._popup = Popup(title="Selecione a língua falada",
                                content=content, size_hint=(0.9, 0.9))
            self._popup.open()

    def save(self, lang):
        if (lang != 'Clique aqui para escolher a língua'):
            self.label1_wid.text = '\n'.join(Root.loadedList)
            Root.choosenLang = lang
            self.dismiss_popup()

    def convert(self):
        if (Root.choosenLang != '' and len(Root.loadedList) > 0):
            while (Root.loadedList):
                mime = Root.loadedList[0].split('.')
                text = ''
                if (mime[len(mime) - 1] == 'pdf'):
                    raw = parser.from_file(Root.fullpath[0])
                    content = raw['content'].split('\n')
                    i = 0
                    while i < len(content):
                        if content[i] == '' or content[i] == ' ':
                            del content[i]
                        else:
                            i += 1
                    text = '\n'.join(content)

                elif (mime[len(mime) - 1] == 'docx'):
                    doc = Document(Root.fullpath[0])
                    content = []
                    for i in range(len(doc.paragraphs)):
                        if (doc.paragraphs[i].text != '' and doc.paragraphs[i].text != ' '):
                            content.append(doc.paragraphs[i].text)
                    text = '\n'.join(content)
                else:
                    content = []
                    with open(Root.fullpath[0], 'r', encoding='utf-8') as f:
                        for line in f:
                            content.append(line.split('\n')[0])
                            if(content[0] != '' and content[0] != ' '):
                                text += line
                            del content[0]
                    f.close()
                self.use_gtts(text, Root.fullpath[0], mime[len(mime) - 1])
                del Root.loadedList[0]
                del Root.fullpath[0]
                self.label1_wid.text = '\n'.join(Root.loadedList)

        elif (len(Root.loadedList) > 0):
            self.label1_wid.text = 'Selecione a língua a ser falada no áudio'
        else:
            self.label1_wid.text = 'Não há arquivos na fila'

    def use_gtts(self, text, path, type):
        try:
            lang = Root.choosenLang.split(':')[0]
            name = ''
            if (type == 'docx'):
                name = path[:-5]
            else:
                name = path[:-4]

            name += '.mp3'
            i = 1

            while(os.path.exists(name)):
                name = path[:-3]
                name += '(' + str(i) + ').mp3'
                i += 1

            output = gTTS(text=text, lang=lang, slow=False)
            output.save(name)

            Root.finishedList.append(Root.loadedList[0])

            self.label3_wid.text = '\n'.join(Root.finishedList)

        except:
            path = path.split('\\')
            Root.failedList.append(Root.loadedList[0])
            self.label4_wid.text = '\n'.join(Root.failedList)


class MP3Conversion(App):
    pass


Factory.register('Root', cls=Root)
Factory.register('LoadDialog', cls=LoadDialog)
Factory.register('LangDialog', cls=LangDialog)

if __name__ == '__main__':
    MP3Conversion().run()
